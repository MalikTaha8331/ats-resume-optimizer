"""
cv_extractor.py
Extracts raw text from an uploaded CV file (PDF, DOCX, or plain text/pasted).
This is step 1 of the pipeline: get clean raw text, preserve structure as much as possible.
"""

import os
import pdfplumber
from docx import Document


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF, page by page, preserving line breaks."""
    full_text = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                full_text.append(text)
    return "\n".join(full_text)


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file, paragraph by paragraph."""
    doc = Document(filepath)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also grab text inside tables (some CVs use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)


def extract_cv_text(filepath: str = None, raw_text: str = None) -> dict:
    """
    Main entry point. Accepts either a file path (PDF/DOCX) or raw pasted text.
    Returns a dict with the extracted text and metadata about extraction quality,
    since ATS-parsability of the ORIGINAL file is itself a signal we score later.
    """
    if raw_text:
        return {
            "text": raw_text.strip(),
            "source_type": "pasted_text",
            "extraction_warnings": []
        }

    if not filepath or not os.path.exists(filepath):
        raise FileNotFoundError(f"CV file not found: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()
    warnings = []

    if ext == ".pdf":
        text = extract_text_from_pdf(filepath)
        if not text.strip():
            warnings.append(
                "No extractable text found in PDF. This likely means the CV is "
                "image-based (scanned) or uses non-standard fonts/encoding. "
                "Real ATS systems will also fail to read this CV — this is a critical finding."
            )
        source_type = "pdf"

    elif ext == ".docx":
        text = extract_text_from_docx(filepath)
        source_type = "docx"

    elif ext == ".txt":
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        source_type = "txt"

    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")

    return {
        "text": text.strip(),
        "source_type": source_type,
        "extraction_warnings": warnings
    }


if __name__ == "__main__":
    # quick manual test
    sample = "John Doe\nSoftware Engineer\nSkills: Python, Flask, SQL"
    result = extract_cv_text(raw_text=sample)
    print(result)
